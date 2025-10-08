
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Mnemonic/Seed Phrase Finder — Pro Edition (PySide6)
===================================================
РУССКИЙ ИНТЕРФЕЙС. Современная тёмная тема.

Что умеет:
- Надёжный поиск реальных BIP-39 seed-фраз (12/15/24 слов) без "генерации" из разрозненных слов.
- Понимает разные способы записи: в строку, в столбик, по 3–4 слова в строке, с номерами/точками/прочими разделителями.
- Учитывает списки вида "1. слово", "2) слово", маркированные списки (•, -, *), смешанные цифры/точки/скобки.
- Сканирует файлы нужных расширений (задаёте сами): .txt, .log, .md, .csv, .html/.htm, .docx, .doc (Windows COM), .xlsx, .xls, + опционально .pdf и .rtf (если библиотеки установлены).
- Результаты группируются: валидные по checksum и «почти» (невалидные) — всё без дублей. Экспорт в .txt и .csv.
- Быстро работает: многопроцессорный скан, отмена/пауза, прогресс, двойной клик — открыть файл.

Важно:
- Мы НИКОГДА не переставляем и не комбинируем слова. Берём только непрерывные последовательности в исходном порядке,
  разрешая между словами только «шум» (цифры, пунктуация, маркеры списка, разрывы строк). Любые другие буквеные слова
  разрывают последовательность. Так мы исключаем генерацию фраз «из воздуха».
"""
from __future__ import annotations

import csv
import os
import string
import re
import sys
import tempfile
from dataclasses import dataclass
from datetime import datetime
from multiprocessing import Pool, cpu_count
from pathlib import Path
from typing import Iterable, List, Tuple, Dict, Set

# ---- Third‑party libs ----
# pip install mnemonic PySide6 python-docx openpyxl xlrd pdfminer.six striprtf beautifulsoup4 lxml pywin32
from PySide6.QtCore import QThread, Signal, QUrl, Qt
from PySide6.QtGui import QDesktopServices, QIcon, QAction
from PySide6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton,
    QCheckBox, QListWidget, QProgressBar, QFileDialog, QMessageBox, QListWidgetItem,
    QSplitter, QPlainTextEdit, QMenuBar, QStatusBar, QAbstractItemView
)
from mnemonic import Mnemonic


# ---- Global exception hook ----
def _qt_fatal(msg: str):
    try:
        with open(f"runtime_error_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log", "w", encoding="utf-8") as f:
            f.write(msg)
    except Exception:
        pass
    try:
        QMessageBox.critical(None, "Mnemonic Finder — Ошибка", msg)
    except Exception:
        pass

def _excepthook(etype, value, tb):
    import traceback
    txt = "".join(traceback.format_exception(etype, value, tb))
    _qt_fatal(txt)

sys.excepthook = _excepthook
# -------------------- Constants --------------------
VALID_LENGTHS = (12, 15, 18, 24)
# Глобальный флаг строгих разрывов (наследуется из GUI через переменную окружения)
try:
    STRICT = (os.environ.get("MF_STRICT_BREAKS", "0") == "1")
except Exception:
    STRICT = False

BIP39 = Mnemonic("english")
BIPSET = set(BIP39.wordlist)

# --- ETH helpers: detect private keys (64 hex) and addresses (0x40 hex). EIP-55 check optional ---
ETH_PRIV_RE = re.compile(r"\b(?:0x)?[0-9A-Fa-f]{64}\b")
ETH_ADDR_RE = re.compile(r"\b0x[0-9A-Fa-f]{40}\b")

def _keccak256_hex(s: str) -> str | None:
    try:
        import sha3  # pysha3
        k = sha3.keccak_256()
        k.update(s.encode('ascii'))
        return k.hexdigest()
    except Exception:
        return None
        return None
        return None

def _to_eip55(addr_noprefix: str) -> str | None:
    h = _keccak256_hex(addr_noprefix.lower())
    if not h:
        return None
    out = []
    for ch, hh in zip(addr_noprefix.lower(), h):
        if ch in '0123456789':
            out.append(ch)
        else:
            out.append(ch.upper() if int(hh, 16) >= 8 else ch)
    return ''.join(out)

def find_eth_keys(text: str):  # -> list[tuple[str,str]]
    found = []
    seen = set()
    for m in ETH_PRIV_RE.finditer(text):
        val = m.group(0)
        val_norm = val[2:] if val.lower().startswith('0x') else val
        if len(val_norm) == 64:
            kv = ('PRIV', val)
            if kv not in seen:
                seen.add(kv); found.append(kv)
    for m in ETH_ADDR_RE.finditer(text):
        addr = m.group(0)
        hexpart = addr[2:]
        if hexpart.islower() or hexpart.isupper():
            kv = ('ADDR', addr)
        else:
            cs = _to_eip55(hexpart)
            kv = ('ADDR', addr) if (cs is None or cs == hexpart) else ('ADDR_BADCHECK', addr)
        if kv not in seen:
            seen.add(kv); found.append(kv)
    return found


# Разрешаем только эти «шумовые» токены между словами сид-фразы
# Цифры, пунктуация, обрамляющие скобки/кавычки, маркеры списка, дефисы и т.п.
NOISE_PATTERN = r"""(?x)
    (?: ^\s*$ )                                  # пустая строка
  | (?: ^\s*[\-\*\u2022\u2023\u25E6\u2043]\s*$ ) # маркер: -, *, • и т.п.
"""

# Регэкспы
WORD_RE = re.compile(r"[A-Za-z]+", re.IGNORECASE)
NUM_ITEM_RE = re.compile(r"^\s*(\d+)[\.\)\-:]\s*(.*)$")  # "1. word" / "1) word" / "1-word"
BULLET_ITEM_RE = re.compile(r"^\s*[\-\*\u2022\u2023\u25E6\u2043]\s*(.+)$")  # "- word" / "• word"

# -------------------- Text extraction --------------------
def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    try:
        if ext == ".docx":
            from docx import Document
            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs]
            # Таблицы
            for t in doc.tables:
                for row in t.rows:
                    parts.append(" ".join(cell.text for cell in row.cells))
            return "\n".join(parts)

        if ext == ".doc":
            # Windows only
            import win32com.client  # type: ignore
            word = win32com.client.Dispatch('Word.Application')
            doc = word.Documents.Open(str(path), ReadOnly=True)
            tmp = Path(tempfile.gettempdir()) / (path.stem + ".txt")
            doc.SaveAs(str(tmp), FileFormat=2)
            doc.Close(False); word.Quit()
            return tmp.read_text(encoding='utf-8', errors='ignore')

        if ext == ".xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            rows = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    rows.append(" ".join(str(c) for c in row if c))
            return "\n".join(rows)

        if ext == ".xls":
            import xlrd
            wb = xlrd.open_workbook(str(path))
            rows = []
            for sh in wb.sheets():
                for r in range(sh.nrows):
                    rows.append(" ".join(str(sh.cell_value(r, c)) for c in range(sh.ncols) if sh.cell_value(r, c)))
            return "\n".join(rows)

        if ext in (".html", ".htm"):
            try:
                from bs4 import BeautifulSoup  # type: ignore
                raw = path.read_text(encoding='utf-8', errors='ignore')
                soup = BeautifulSoup(raw, "lxml")
                # убирать скрипты/стили
                for bad in soup(["script", "style", "noscript"]): bad.decompose()
                return soup.get_text(separator="\n")
            except Exception:
                raw = path.read_text(encoding='utf-8', errors='ignore')
                raw = re.sub(r'<script.*?>.*?</script>', '', raw, flags=re.S | re.I)
                raw = re.sub(r'<style.*?>.*?</style>', '', raw, flags=re.S | re.I)
                text = re.sub(r'<[^>]+>', ' ', raw)
                return text

        if ext == ".pdf":
            try:
                from pdfminer.high_level import extract_text as pdf_extract  # type: ignore
                return pdf_extract(str(path))
            except Exception:
                return ""

        if ext == ".rtf":
            try:
                from striprtf.striprtf import rtf_to_text  # type: ignore
                raw = path.read_text(encoding='utf-8', errors='ignore')
                return rtf_to_text(raw)
            except Exception:
                return ""

        # Прочее как текст
        data = path.read_bytes()
        try:
            return data.decode('utf-8', errors='ignore')
        except Exception:
            return data.decode('latin-1', errors='ignore')
    except Exception:
        return ""

# -------------------- Tokenization & detection --------------------

# --- NEW: Strict tokenizer for LINEAR mode (break on punctuation and end-of-line; allow numeric labels) ---
LABEL_WORD_RE = re.compile(r'^(?:#?\d{1,2})(?:[).:-])?([A-Za-z]+)$')
LABEL_ONLY_RE = re.compile(r'^#?\d{1,2}[).:-]?$', re.IGNORECASE)

def tokenize_with_breaks_strict(text: str) -> List[str]:
    '''
    Tokenize text so that any non-letter punctuation acts as a HARD break between BIP words.
    Recognizes numeric labels like '1)word', '#2:word', '03.word' and extracts 'word' without breaking.
    Adds a BREAK at end of each line.
    '''
    tokens: List[str] = []
    for line in text.splitlines():
        parts = re.split(r'[,\s]+', line)
        for raw in parts:
            if not raw:
                continue
            raw = raw.strip()
            m = LABEL_WORD_RE.match(raw)
            if m:
                tokens.append(m.group(1).lower());  # keep as word
                continue
            if WORD_RE.fullmatch(raw):
                tokens.append(raw.lower());  # word
                continue
            if LABEL_ONLY_RE.match(raw):
                # pure numeric label like "1)", "#2." -> ignore (but do NOT break)
                continue
            # anything else (e.g., urls, slashes, pipes, etc.) => HARD break
            tokens.append("|BREAK|")
        tokens.append("|BREAK|")  # end-of-line is also a break
    # collapse multiple BREAKs
    out: List[str] = []
    for t in tokens:
        if t == "|BREAK|":
            if out and out[-1] == "|BREAK|":
                continue
        out.append(t)
    return out


def tokenize_with_breaks_strict(text: str) -> List[str]:
    '''
    Tokenize text so that any non-letter punctuation acts as a HARD break between BIP words.
    Recognizes numeric labels like '1)word', '#2:word', '03.word' and extracts 'word' without breaking.
    Adds a BREAK at end of each line.
    Memory‑friendly: avoids building long runs of repeated BREAK tokens.
    '''
    tokens: List[str] = []
    last_was_break = True  # prevent leading BREAK

    for line in text.splitlines():
        parts = re.split(r'[,\s]+', line)
        for raw in parts:
            if not raw:
                continue
            raw = raw.strip()
            m = LABEL_WORD_RE.match(raw)
            if m:
                tokens.append(m.group(1).lower())
                last_was_break = False
                continue
            if WORD_RE.fullmatch(raw):
                tokens.append(raw.lower())
                last_was_break = False
                continue
            if LABEL_ONLY_RE.match(raw):
                # numeric label like "1)", "#2." -> ignore (no break)
                continue
            # anything else => HARD break (but don't append duplicates)
            if not last_was_break:
                tokens.append("|BREAK|")
                last_was_break = True
        # end-of-line is a break (avoid duplicates)
        if not last_was_break:
            tokens.append("|BREAK|")
            last_was_break = True
    return tokens


def tokenize_with_breaks_normal(text: str) -> List[str]:
    """
    «Мягкий» токенайзер: пока просто совпадает со строгим, чтобы
    логика оставалась простой и без генерации. При желании можно
    ослабить правила, но это не требуется после наших правок.
    """
    return tokenize_with_breaks_strict(text)


# --- NEW: tiny Levenshtein within 1 edit, for speed ---
def within_one_edit(a: str, b: str) -> bool:
    if a == b: return True
    la, lb = len(a), len(b)
    if abs(la - lb) > 1: return False
    # ensure a is shorter or equal
    if la > lb:
        a, b = b, a
        la, lb = lb, la
    i = j = diffs = 0
    while i < la and j < lb:
        if a[i] == b[j]:
            i += 1; j += 1
        else:
            diffs += 1
            if diffs > 1: return False
            if la == lb:
                i += 1; j += 1  # substitution
            else:
                j += 1  # insertion in longer word
    # tail
    if j < lb or i < la:
        diffs += 1
    return diffs <= 1

def autocorrect_phrase_12_one_typo(phrase: str) -> str | None:
    '''
    Returns corrected 12-word phrase where each word is either in BIPSET
    or within one edit of a BIP word. If any word cannot be corrected within 1 edit, returns None.
    '''
    words = phrase.split()
    if len(words) != 12:
        return None
    corrected: List[str] = []
    for w in words:
        if w in BIPSET:
            corrected.append(w); continue
        # try to find a close match
        # Fast check: try words with same first letter to cut search
        candidates = [bw for bw in BIPSET if bw[0:1] == w[0:1] and abs(len(bw)-len(w)) <= 1]
        found = None
        for bw in candidates:
            if within_one_edit(w, bw):
                found = bw; break
        if not found:
            # fallback: scan a bit more broadly but still bounded
            for bw in BIPSET:
                if abs(len(bw)-len(w)) <= 1 and within_one_edit(w, bw):
                    found = bw; break
        if not found:
            return None
        corrected.append(found)
    return " ".join(corrected)

# --- NEW: Electrum v2+ checksum (seed version prefix) ---
ELECTRUM_PREFIXES = ("01","100","101","102")

def _normalize_text_electrum(seed: str) -> str:
    import unicodedata, string as _string
    # mimic electrum.normalize_text (simplified, without CJK whitespace rule)
    seed = unicodedata.normalize('NFKD', seed).lower()
    seed = ''.join(c for c in seed if not unicodedata.combining(c))
    seed = ' '.join(seed.split())
    return seed

def is_electrum_new_seed(mnemonic: str) -> bool:
    '''
    Check Electrum 'new' seed (v2+) by HMAC-SHA512('Seed version', normalized_text) having a known prefix.
    '''
    import hashlib, hmac, binascii
    x = _normalize_text_electrum(mnemonic)
    h = hmac.new(b"Seed version", x.encode('utf-8'), hashlib.sha512).hexdigest()
    return any(h.startswith(pfx) for pfx in ELECTRUM_PREFIXES)

def words_in_line(line: str) -> List[str]:
    """Все английские слова в нижнем регистре из строки."""
    return [w.lower() for w in WORD_RE.findall(line)]

def bip_words_in_line(line: str) -> List[str]:
    """Только слова из BIP-39 в нижнем регистре."""
    return [w for w in words_in_line(line) if w in BIPSET]

def collapse_inline_noise(tokens: List[str]) -> List[str]:
    """
    Из последовательности токенов берём только подряд идущие BIP-39 слова.
    Между ними разрешаем ТОЛЬКО: цифры, пунктуацию, маркеры списков и пустые строки.
    Любые другие буквеные слова обрывают последовательность.
    """
    result: List[str] = []
    for t in tokens:
        if t in BIPSET:
            result.append(t)
        else:
            # t содержит буквы? — обрываем текущую цепочку (вставляем разделитель)
            if re.search(r"[A-Za-z]", t):
                result.append("|BREAK|")
            else:
                # это «шум» — игнорируем
                continue
    # сжать множественные BREAK
    squeezed: List[str] = []
    for t in result:
        if t == "|BREAK|":
            if squeezed and squeezed[-1] != "|BREAK|":
                squeezed.append(t)
        else:
            squeezed.append(t)
    return squeezed


def segments_from_stream(stream: List[str]) -> Iterable[List[str]]:
    """
    Возвращает непрерывные сегменты BIP-слов, разделённые |BREAK|.
    Никаких скользящих окон — только цельные сегменты.
    """
    segment: List[str] = []
    for t in stream + ["|BREAK|"]:
        if t == "|BREAK|":
            if segment:
                yield segment
            segment = []
        else:
            segment.append(t)
def sliding_phrases_from_stream(stream: List[str]) -> Iterable[List[str]]:
    """
    Берём скользящим окном 12/15/24 по непрерывным участкам (без BREAK).
    """
    segment: List[str] = []
    for t in stream + ["|BREAK|"]:
        if t == "|BREAK|":
            if len(segment) >= 12:
                for k in VALID_LENGTHS:
                    if len(segment) >= k:
                        for i in range(0, len(segment) - k + 1):
                            yield segment[i:i+k]
            segment = []
        else:
            segment.append(t)

def list_mode_candidates(lines: List[str]) -> Iterable[List[str]]:
    """
    Режим «в столбик»: подряд идущие строки, где у каждой — ровно одно BIP-слово
    (может быть префикс нумерации/маркера).
    """
    buf: List[str] = []
    def flush():
        nonlocal buf
        if len(buf) >= 12:
            for k in VALID_LENGTHS:
                if len(buf) >= k:
                    for i in range(0, len(buf)-k+1):
                        yield buf[i:i+k]
        buf = []

    for line in lines + [""]:
        line_stripped = line.strip()
        if not line_stripped or re.match(NOISE_PATTERN, line_stripped):
            # пустая/шумовая строка — фрагмент закончился
            yield from flush()
            continue

        m = NUM_ITEM_RE.match(line_stripped)
        if m:
            rest = m.group(2)
            bwords = bip_words_in_line(rest)
        else:
            m2 = BULLET_ITEM_RE.match(line_stripped)
            if m2:
                bwords = bip_words_in_line(m2.group(1))
            else:
                bwords = bip_words_in_line(line_stripped)

        if len(bwords) == 1:
            buf.append(bwords[0])
            continue
        elif len(bwords) in (3, 4):
            # поддержка по 3–4 слова в строке
            buf.extend(bwords)
            continue
        else:
            # строка не подходит — закрываем текущий блок
            yield from flush()

    # финальный flush на случай, если закончилось без шума
    yield from flush()

def find_phrases_robust(text: str) -> Tuple[List[str], List[str]]:
    """
    Возвращает два списка:
      valid_phrases — с корректной checksum,
      near_phrases  — из слов BIP-39, но не проходящие checksum (для справки).
    Детектор состоит из двух независимых режимов: LINEAR и LIST.
    """
    lines = text.splitlines()

    # 12-словники, только если вся строка = 12 BIP-слов
    line_12_candidates: List[str] = []
    if os.environ.get('STRICT_12_LIST_ONLY','0') == '1':
        for _line in lines:
            parts = re.split(r'[,\s]+', _line)
            _words: List[str] = []
            _ok = True
            for _raw in parts:
                if not _raw: continue
                _raw = _raw.strip()
                m = LABEL_WORD_RE.match(_raw)
                if m:
                    _words.append(m.group(1).lower()); continue
                if LABEL_ONLY_RE.match(_raw):
                    continue
                w = _raw.strip(string.punctuation).lower()
                if w and w in BIPSET:
                    _words.append(w)
                else:
                    _ok = False; break
            if _ok and len(_words) == 12:
                line_12_candidates.append(' '.join(_words))
        # --- two-line join for 12 words (strict anti-false mode) ---
        # --- multi-line patterns 12x1 and 6x2 (strict anti-false mode) ---
        # Pre-parse each line into BIP-only word list (labels ignored); None if any non-BIP token present.
        parsed_lines: List[List[str] | None] = []
        for _line in lines:
            parts = re.split(r'[,\s]+', _line)
            words_line: List[str] = []
            ok_line = True
            for _raw in parts:
                if not _raw: 
                    continue
                _raw = _raw.strip()
                m = LABEL_WORD_RE.match(_raw)
                if m:
                    words_line.append(m.group(1).lower()); continue
                if LABEL_ONLY_RE.match(_raw):
                    continue
                w = _raw.strip(string.punctuation).lower()
                if w and w in BIPSET:
                    words_line.append(w)
                else:
                    ok_line = False; break
            parsed_lines.append(words_line if ok_line else None)

        # 12 consecutive lines with exactly 1 word each
        for i in range(0, max(0, len(parsed_lines) - 12 + 1)):
            block = parsed_lines[i:i+12]
            if all(bl is not None and len(bl) == 1 for bl in block):
                phrase = ' '.join(w for bl in block for w in bl)
                if phrase not in line_12_candidates:
                    line_12_candidates.append(phrase)

        # 6 consecutive lines with exactly 2 words each
        for i in range(0, max(0, len(parsed_lines) - 6 + 1)):
            block = parsed_lines[i:i+6]
            if all(bl is not None and len(bl) == 2 for bl in block):
                phrase = ' '.join(w for bl in block for w in bl)
                if phrase not in line_12_candidates:
                    line_12_candidates.append(phrase)
    
        for i in range(len(lines)-1):
            words: List[str] = []
            ok = True
            for _line in (lines[i], lines[i+1]):
                parts = re.split(r'[,\s]+', _line)
                for _raw in parts:
                    if not _raw:
                        continue
                    _raw = _raw.strip()
                    m = LABEL_WORD_RE.match(_raw)
                    if m:
                        words.append(m.group(1).lower()); continue
                    if LABEL_ONLY_RE.match(_raw):
                        continue
                    w = _raw.strip(string.punctuation).lower()
                    if w and w in BIPSET:
                        words.append(w)
                    else:
                        ok = False; break
                if not ok:
                    break
            if ok and len(words) == 12:
                phrase = ' '.join(words)
                if phrase not in line_12_candidates:
                    line_12_candidates.append(phrase)
    # подготовим поток токенов (строим только один, чтобы не удваивать память)
    streams = tokenize_with_breaks_strict(text) if STRICT else tokenize_with_breaks_normal(text)
    linear_candidates: List[str] = []
    for seg in segments_from_stream(streams):
        if len(seg) in VALID_LENGTHS:
            if os.environ.get("STRICT_12_LIST_ONLY","0") == "1" and len(seg) == 12:
                continue
            linear_candidates.append(" ".join(seg))
    # --- LIST MODE ---
    if os.environ.get("STRICT_12_LIST_ONLY","0") == "1":
        list_candidates = []
    else:
        list_candidates = [" ".join(cand) for cand in list_mode_candidates(lines)]

    # Объединяем кандидатов, проверяем checksum
    # Приоритет по длине: 24 → 15 → 12, и без дублей
    candidates = []
    seen: Set[str] = set()
    for c in sorted(linear_candidates + list_candidates + line_12_candidates, key=lambda s: -len(s.split())):
        if c not in seen:
            candidates.append(c); seen.add(c)

    valid, near = [], []
    for phrase in candidates:
        words = phrase.split()
        if len(words) not in VALID_LENGTHS:
            continue
        # Pure BIP39 path
        if all((w in BIPSET) for w in words):
            if BIP39.check(phrase):
                valid.append(phrase)
            else:
                near.append(phrase)
            continue
        # Electrum path: allow 12/24 any english words (not only BIP39)
        if all(re.fullmatch(r"[A-Za-z]+", w) for w in words) and is_electrum_new_seed(phrase):
            near.append(phrase)

    return valid, near
def is_suspicious(phrase: str) -> bool:
    """
    Определяет «подозрительность» валидной фразы: повторяются 2–3 слова, низкая уникальность,
    повторяющиеся биграммы. Такие фразы убираем из «валидных» и заносим в отдельный список/файл.
    """
    words = phrase.split()
    n = len(words)
    if n == 0:
        return False
    # Частоты слов
    from collections import Counter
    cnt = Counter(words)
    uniq = len(cnt)
    max_freq = max(cnt.values())
    unique_ratio = uniq / n
    max_rel = max_freq / n

    # Повторяющиеся биграммы
    bigrams = list(zip(words, words[1:]))
    if bigrams:
        bcnt = Counter(bigrams)
        rep_bigrams = sum(1 for b,c in bcnt.items() if c >= 2)
        bigram_rep_ratio = rep_bigrams / max(1, len(bigrams))
    else:
        bigram_rep_ratio = 0.0

    # Хард правила (консервативно):
    #  - очень мало разных слов (<= n/3) или всего <= 5 уникальных
    #  - одно слово встречается >= 1/3 всех позиций
    #  - повторяющиеся биграммы занимают >= 40% всех биграмм
    if uniq <= max(5, n // 3):
        return True
    if max_rel >= 1/3:
        return True
    if bigram_rep_ratio >= 0.4:
        return True

    return False


# -------------------- Worker --------------------
def scan_file(path_str: str) -> Tuple[str, List[str], List[str], List[str], List[tuple[str,str]]]:
    text = extract_text(Path(path_str))
    if not text:
        return path_str, [], [], [], []
    valid, near = find_phrases_robust(text)
    suspicious = [v for v in valid if is_suspicious(v)]
    valid_clean = [v for v in valid if v not in suspicious]
    eth = find_eth_keys(text) if os.environ.get('FIND_ETH','0') == '1' else []
    return path_str, valid_clean, near, suspicious, eth

class Worker(QThread):
    progress = Signal(int, int)
    found_valid = Signal(str, str)    # phrase, file
    found_near = Signal(str, str)     # phrase, file
    found_suspicious = Signal(str, str)  # phrase, file
    found_eth = Signal(str, str)  # kind:value, file
    done = Signal()

    def __init__(self, root: Path, exts: Set[str], exclude: Set[str], max_mb: int, follow_hidden: bool):
        super().__init__()
        self.root = root
        self.exts = {e if e.startswith('.') else f'.{e}' for e in exts}
        self.exclude = {x.lower() for x in exclude}
        self.max_bytes = max_mb * 1024 * 1024 if max_mb > 0 else None
        self.follow_hidden = follow_hidden
        self._pool = None

    def run(self):
        files: List[str] = []
        for p in self.root.rglob('*'):
            if not p.is_file(): continue
            if p.suffix.lower() not in self.exts: continue
            if any(seg.lower() in self.exclude for seg in p.parts): continue
            if not self.follow_hidden and any(part.startswith('.') for part in p.parts): continue
            if self.max_bytes and p.stat().st_size > self.max_bytes: continue
            files.append(str(p))

        total = len(files)
        procs = max(cpu_count() - 1, 1)
        count = 0
        with Pool(procs) as pool:
            for path, valid, near, suspicious, eth in pool.imap_unordered(scan_file, files):
                count += 1
                for phr in valid:
                    self.found_valid.emit(phr, path)
                for phr in near:
                        self.found_near.emit(phr, path)
                for phr in suspicious:
                    self.found_suspicious.emit(phr, path)
                for k, v in eth:
                    self.found_eth.emit(f"{k}:{v}", path)
                self.progress.emit(count, total)
        self.done.emit()

# -------------------- GUI --------------------
class Main(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Mnemonic Finder — Pro")
        self.setWindowIcon(QIcon.fromTheme("preferences-system-privacy"))
        self.resize(1000, 700)
        self._apply_dark_qss()

        # Top bar
        self.dir_edit = QLineEdit(); self.dir_edit.setPlaceholderText("Выберите папку для сканирования…")
        self.btn_browse = QPushButton("Обзор")
        self.btn_browse.clicked.connect(self.browse_folder)

        self.ext_edit = QLineEdit(".txt,.log,.md,.csv,.html,.htm,.doc,.docx,.xls,.xlsx,.pdf,.rtf")
        self.ext_edit
        self.ext_edit.setMinimumWidth(300)
        self.ext_edit.setToolTip("Через запятую. Точки можно опускать: txt,log,docx")

        self.exclude_edit = QLineEdit("node_modules,.git,.venv,venv,__pycache__")
        self.exclude_edit.setToolTip("Папки-исключения через запятую")

        self.size_edit = QLineEdit("50"); self.size_edit.setFixedWidth(60)
        self.chk_hidden = QCheckBox("Скрытые")
        self.chk_strict = QCheckBox("Меньше мусора (Linear)")
        self.chk_strict.setChecked(True)
        self.chk_typo12 = QCheckBox("Допускать 1 ошибку")
        self.chk_near = QCheckBox("Показывать невалидные (near)")
        self.chk_12list = QCheckBox("12 слов — только как список (anti-false)")
        self.chk_eth = QCheckBox("Искать ETH ключи/адреса")

        self.btn_start = QPushButton("Старт"); self.btn_start.setProperty("accent", True)
        self.btn_stop = QPushButton("Стоп"); self.btn_stop.setEnabled(False)
        self.btn_export = QPushButton("Экспорт CSV"); self.btn_export.setEnabled(False)

        # Lists
        self.list_valid = QListWidget()
        self.list_near = QListWidget()
        self.list_susp = QListWidget()
        self.list_valid.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.list_near.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.list_susp.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.list_valid.itemDoubleClicked.connect(self._reveal_in_folder)
        self.list_near.itemDoubleClicked.connect(self._reveal_in_folder)
        self.list_susp.itemDoubleClicked.connect(self._reveal_in_folder)

        # Log/preview
        self.log = QPlainTextEdit(); self.log.setReadOnly(True); self.log.setTextInteractionFlags(Qt.TextSelectableByMouse | Qt.TextSelectableByKeyboard); self.log.setPlaceholderText("Журнал работы…")

        # Layouts
        top1 = QHBoxLayout()
        top1.addWidget(QLabel("Папка:")); top1.addWidget(self.dir_edit); top1.addWidget(self.btn_browse)
        top2 = QHBoxLayout()
        top2b = QHBoxLayout()
        top2.addWidget(QLabel("Расширения:")); top2.addWidget(self.ext_edit, 2)
        top2.addWidget(QLabel("Искл.:")); top2.addWidget(self.exclude_edit, 2)
        top2.addWidget(QLabel("Макс, МБ:")); top2.addWidget(self.size_edit)
        top2.addWidget(self.chk_hidden)
        top2.addWidget(self.btn_start); top2.addWidget(self.btn_stop); top2.addWidget(self.btn_export)

        split = QSplitter()
        left = QVBoxLayout(); lw = QWidget(); lw.setLayout(left)
        left.addWidget(QLabel("Валидные (checksum ОК)"))
        left.addWidget(self.list_valid)
        left.addWidget(QLabel("Подозрительные (валидны, но низкая уникальность/повторы)"))
        left.addWidget(self.list_susp)
        left.addWidget(QLabel("Лог Electrum"))
        self.log_electrum = QPlainTextEdit()
        self.log_electrum.setReadOnly(True)
        self.log_electrum.setTextInteractionFlags(Qt.TextSelectableByMouse | Qt.TextSelectableByKeyboard)
        left.addWidget(self.log_electrum)

        
        top2.addWidget(self.chk_strict)
        top2.addWidget(self.chk_typo12)
        top2b.addWidget(self.chk_near)
        top2b.addWidget(self.chk_12list)
        top2b.addWidget(self.chk_eth)
        right = QVBoxLayout(); rw = QWidget(); rw.setLayout(right)
        right.addWidget(QLabel("Журнал"))
        right.addWidget(self.log)

        split.addWidget(lw); split.addWidget(rw); split.setStretchFactor(0, 3); split.setStretchFactor(1, 2)

        self.progress_bar = QProgressBar()
        self.status = QStatusBar()
        self.lbl_counts = QLabel("")
        self.lbl_counts.setTextFormat(Qt.RichText)
        self.status.addPermanentWidget(self.lbl_counts)

        root = QVBoxLayout(self)
        root.addLayout(top1); root.addLayout(top2); root.addLayout(top2b); root.addWidget(split); root.addWidget(self.progress_bar); root.addWidget(self.status)

        # State
        self.worker: Worker | None = None
        self.results_dir: Path | None = None
        self.timestamp: str | None = None
        self.seen_valid: Dict[int, Set[str]] = {12:set(), 15:set(), 18:set(), 24:set()}
        self.seen_near: Set[str] = set()
        self.seen_susp: Set[str] = set()
        self.count_seed: Dict[int,int] = {12:0,15:0,18:0,24:0}
        self.count_electrum: int = 0
        self.count_eth: int = 0

        # Signals
        self.btn_start.clicked.connect(self.start_scan)
        self.btn_stop.clicked.connect(self.stop_scan)
        self.btn_export.clicked.connect(self.export_csv)

    # --------- UI helpers ---------
    def _apply_dark_qss(self):
        self.setStyleSheet("""
        QWidget { background-color: #131416; color: #E8EAED; font: 14px 'Segoe UI'; }
        QLabel { color: #AEB4BE; }
        QLineEdit, QPlainTextEdit, QListWidget {
            background-color: #1A1C1F; border: 1px solid #2A2E33; border-radius: 6px; padding: 6px;
        }
        QPushButton {
            background-color: #24272C; border: 1px solid #2F353C; padding: 8px 14px; border-radius: 8px;
        }
        QPushButton:hover { border-color: #3C444D; }
        QPushButton[accent="true"] { background-color: #2D7DFF; color: white; border: none; }
        QPushButton[accent="true"]:disabled { background-color: #2D7DFF55; }
        QProgressBar { background: #1A1C1F; border: 1px solid #2A2E33; border-radius: 6px; text-align: center; }
        QProgressBar::chunk { background: #2D7DFF; }
        QSplitter::handle { background: #1A1C1F; }
        QListWidget::item { padding: 4px; }
        """)

    def browse_folder(self):
        d = QFileDialog.getExistingDirectory(self, "Выбор папки")
        if d: self.dir_edit.setText(d)

    def _parse_exts(self) -> Set[str]:
        raw = [e.strip().lower() for e in self.ext_edit.text().replace(" ", "").split(",") if e.strip()]
        return {e if e.startswith('.') else f'.{e}' for e in raw}

    def _parse_exclude(self) -> Set[str]:
        return {x.strip() for x in self.exclude_edit.text().split(",") if x.strip()}

    def start_scan(self):
        folder = self.dir_edit.text().strip()
        if not Path(folder).is_dir():
            QMessageBox.warning(self, "Ошибка", "Нужно выбрать существующую папку"); return

        try:
            max_mb = int(self.size_edit.text().strip())
        except Exception:
            max_mb = 50

        self.list_valid.clear(); self.list_near.clear(); self.list_susp.clear(); self.log.clear()
        self.progress_bar.setValue(0); self.progress_bar.setMaximum(0)
        self.btn_start.setEnabled(False); self.btn_stop.setEnabled(True); self.btn_export.setEnabled(False)
        self.seen_valid = {12:set(), 15:set(), 18:set(), 24:set()}
        self.seen_near = set()
        self.seen_susp = set()
        self.seen_electrum_12 = set()
        self.seen_electrum_24 = set()
        self.seen_valid_12_typo = set()
        self.count_eth = 0

        # reset additional tracking sets
        import os as _os
        _os.environ['MF_STRICT_BREAKS'] = '1' if self.chk_strict.isChecked() else '0'
        _os.environ['SHOW_NEAR'] = '1' if self.chk_near.isChecked() else '0'
        _os.environ['STRICT_12_LIST_ONLY'] = '1' if self.chk_12list.isChecked() else '0'
        _os.environ['FIND_ETH'] = '1' if self.chk_eth.isChecked() else '0'
        self.timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        self.results_dir = Path(folder) / f"Results_{self.timestamp}"
        self.results_dir.mkdir(exist_ok=True)

        self.worker = Worker(Path(folder), self._parse_exts(), self._parse_exclude(), max_mb, self.chk_hidden.isChecked())
        self.worker.progress.connect(self.on_progress)
        self.worker.found_valid.connect(self.on_found_valid)
        self.worker.found_near.connect(self.on_found_near)
        self.worker.found_suspicious.connect(self.on_found_susp)
        self.worker.found_eth.connect(self.on_found_eth)
        self.worker.done.connect(self.on_done)
        self.worker.start()
        self.log.appendPlainText(f"Начат поиск в: {folder}")

    def stop_scan(self):
        if self.worker:
            self.worker.terminate(); self.worker = None
            self.on_done(cancelled=True)

    def on_progress(self, c: int, t: int):
        self.progress_bar.setMaximum(t if t > 0 else 0)
        self.progress_bar.setValue(c)

    def _add_result(self, lst: QListWidget, phrase: str, path: str) -> None:
        item = QListWidgetItem(f"{phrase}  |  {path}")
        item.setToolTip(path)
        lst.addItem(item)

    def on_found_valid(self, phrase: str, path: str):
        n = len(phrase.split())
        if phrase not in self.seen_valid[n]:
            self.seen_valid[n].add(phrase)
            self._add_result(self.list_valid, phrase, path)
            try:
                self.count_seed[n] += 1
                self._update_counters()
            except Exception:
                pass
            # Append to files
            try:
                with open(self.results_dir / f"valid_{n}.txt", "a", encoding="utf-8") as f:
                    f.write(f"{phrase} | {path}\n")
            except Exception:
                pass
        # Electrum check (save separately) for 12/24 words
        try:
            n = len(phrase.split())
            if n in (12, 24) and is_electrum_new_seed(phrase):
                seen_set = self.seen_electrum_12 if n == 12 else self.seen_electrum_24
                if phrase not in seen_set:
                    seen_set.add(phrase)
                    self.count_electrum += 1
                    self._update_counters()
                    # write to files
                    fname = f"electrum_valid_{n}.txt"
                    with open(self.results_dir / fname, "a", encoding="utf-8") as f:
                        f.write(f"{phrase} | {path}\n")
                    try:
                        self.log_electrum.appendPlainText(f"{phrase}  |  {path}")
                    except Exception:
                        pass
                    try:
                        self.log_electrum.appendPlainText(f"{phrase}  |  {path}")
                    except Exception:
                        pass
        except Exception:
            pass


    def on_found_susp(self, phrase: str, path: str):
        if phrase not in self.seen_susp:
            self.seen_susp.add(phrase)
            self._add_result(self.list_susp, phrase, path)
            try:
                with open(self.results_dir / "suspicious.txt", "a", encoding="utf-8") as f:
                    f.write(f"{phrase} | {path}\n")
            except Exception:
                pass

    def on_found_near(self, phrase: str, path: str):
        # Only add to the Near list if SHOW_NEAR is enabled; still run Electrum checks below.
        if os.environ.get('SHOW_NEAR','0') == '1' and phrase not in self.seen_near:
            self.seen_near.add(phrase)
            self._add_result(self.list_near, phrase, path)
            try:
                with open(self.results_dir / "near.txt", "a", encoding="utf-8") as f:
                    f.write(f"{phrase} | {path}\n")
            except Exception:
                pass
        # If 12-word and typo mode is ON, try autocorrect and recheck BIP39
        try:
            if self.chk_typo12.isChecked() and len(phrase.split()) == 12:
                corrected = autocorrect_phrase_12_one_typo(phrase)
                if corrected and corrected != phrase and corrected not in self.seen_valid[12] and Mnemonic("english").check(corrected):
                    self.on_found_valid(corrected, path)
                    self.seen_valid_12_typo.add(corrected)
        except Exception:
            pass

        # Electrum check (save separately) for 12/24 words (even if not BIP39-valid)
        try:
            n = len(phrase.split())
            if n in (12, 24) and is_electrum_new_seed(phrase):
                seen_set = self.seen_electrum_12 if n == 12 else self.seen_electrum_24
                if phrase not in seen_set:
                    seen_set.add(phrase)
                    fname = f"electrum_valid_{n}.txt"
                    with open(self.results_dir / fname, "a", encoding="utf-8") as f:
                        f.write(f"{phrase} | {path}\n")
                    try:
                        self.log_electrum.appendPlainText(f"{phrase}  |  {path}")
                    except Exception:
                        pass
                    try:
                        self.log_electrum.appendPlainText(f"{phrase}  |  {path}")
                    except Exception:
                        pass
        except Exception:
            pass


    
    def on_found_eth(self, kind_value: str, path: str):
        if not self.results_dir:
            return
        try:
            with open(self.results_dir / "eth_keys_found.txt", "a", encoding="utf-8") as f:
                f.write(f"{kind_value} | {path}\n")
            self.count_eth += 1
            self._update_counters()
        except Exception:
            pass
    
    def _update_counters(self):
        html = (f"Seed 12 word: <b><span style='color:#28a745'>{self.count_seed[12]}</span></b> &nbsp; "
                f"Seed 15 word: <b><span style='color:#17a2b8'>{self.count_seed[15]}</span></b> &nbsp; "
                f"Seed 24 word: <b><span style='color:#dc3545'>{self.count_seed[24]}</span></b> &nbsp; "
                f"Electrum: <b><span style='color:#ff9800'>{self.count_electrum}</span></b> &nbsp; "
                f"ETH key: <b><span style='color:#9c27b0'>{self.count_eth}</span></b>")
        try:
            self.lbl_counts.setText(html)
        except Exception:
            pass
    def _reveal_in_folder(self, item: QListWidgetItem):
        text = item.text()
        path = text.split("|")[-1].strip()
        if Path(path).exists():
            QDesktopServices.openUrl(QUrl.fromLocalFile(str(Path(path).parent)))

    def on_done(self, cancelled: bool = False):
        self.btn_start.setEnabled(True); self.btn_stop.setEnabled(False); self.btn_export.setEnabled(True)
        self.progress_bar.setMaximum(1); self.progress_bar.setValue(1)
        self.status.showMessage("Отменено пользователем" if cancelled else "Готово", 5000)
        self.log.appendPlainText("Поиск остановлен." if cancelled else "Сканирование завершено.")

    def export_csv(self):
        if not self.results_dir: return
        csv_path = self.results_dir / "results.csv"
        try:
            with open(csv_path, "w", newline="", encoding="utf-8") as f:
                w = csv.writer(f)
                w.writerow(["status", "length", "phrase", "file"])
                for n, phrases in self.seen_valid.items():
                    for phr in sorted(phrases):
                        w.writerow(["valid", n, phr, ""])
                for phr in sorted(self.seen_susp):
                    w.writerow(["suspicious", len(phr.split()), phr, ""])
                
                # Electrum valid
                for phr in sorted(self.seen_electrum_12):
                    w.writerow(["electrum", 12, phr, ""])
                for phr in sorted(self.seen_electrum_24):
                    w.writerow(["electrum", 24, phr, ""])
                # Typo-corrected 12-word valids
                for phr in sorted(self.seen_valid_12_typo):
                    w.writerow(["valid_typo12", 12, phr, ""])

                for phr in sorted(self.seen_near):
                    w.writerow(["near", len(phr.split()), phr, ""])
            QMessageBox.information(self, "Экспорт", f"CSV сохранён: {csv_path}")
            QDesktopServices.openUrl(QUrl.fromLocalFile(str(csv_path.parent)))
        except Exception as e:
            QMessageBox.warning(self, "Экспорт", f"Не удалось сохранить CSV:\n{e}")

def main():
    app = QApplication(sys.argv)
    w = Main()
    w.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()
# ---- Fatal error helper (shows dialog even without Qt) ----
def _fatal_startup(msg: str):
    try:
        # Windows MessageBox without Qt
        import ctypes
        ctypes.windll.user32.MessageBoxW(0, msg, "Mnemonic Finder — Ошибка запуска", 0x00000010)  # MB_ICONERROR
    except Exception:
        pass
    try:
        # Write log next to script
        log_name = f"startup_error_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
        with open(log_name, "w", encoding="utf-8") as f:
            f.write(msg)
    except Exception:
        pass
    sys.exit(1)